# -*- coding: utf-8 -*-
"""
PDF 万能转换工具 - 带 GUI 界面
=====================================
依赖库安装（在终端/命令行执行）：
    pip install PyMuPDF python-docx Pillow

说明：
- PyMuPDF (fitz) 用于 PDF 核心操作（转Word文字、转图片、合并、拆分、加密/解密）
- python-docx 用于生成 Word (.docx) 文件
- Pillow 用于图片格式转换处理
- GUI 使用 Python 自带 tkinter，无需额外安装
- 图片转PDF 使用 fitz 内置功能，无需 poppler

运行方式：
    python pdf_tool.py

Author: zj126-creator
Date: 2026-07-17
"""

import os
import sys
import threading
import traceback
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# ============================================================
#  核心转换逻辑
# ============================================================

def get_output_dir(filepath: str) -> str:
    """在原文件所在目录下创建 Converted 子文件夹"""
    parent = os.path.dirname(os.path.abspath(filepath))
    out_dir = os.path.join(parent, "Converted")
    os.makedirs(out_dir, exist_ok=True)
    return out_dir


def pdf_to_word(pdf_path: str, progress_cb=None) -> str:
    """PDF 转 Word（提取文字生成 .docx）"""
    out_dir = get_output_dir(pdf_path)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    docx_path = os.path.join(out_dir, f"{base_name}.docx")

    doc = Document()
    doc.add_heading(base_name, level=0)

    pdf_doc = fitz.open(pdf_path)
    total = len(pdf_doc)

    for i, page in enumerate(pdf_doc):
        text = page.get_text("text")
        if text.strip():
            doc.add_heading(f"第 {i + 1} 页", level=1)
            doc.add_paragraph(text)
        else:
            doc.add_heading(f"第 {i + 1} 页", level=1)
            doc.add_paragraph("（此页无可提取的文字内容，可能为图片页）")

        if progress_cb:
            progress_cb(i + 1, total)

    doc.save(docx_path)
    pdf_doc.close()
    return docx_path


def pdf_to_images(pdf_path: str, fmt: str = "PNG", dpi: int = 200, progress_cb=None) -> list:
    """PDF 转图片（每页转为高清 PNG/JPG）"""
    out_dir = get_output_dir(pdf_path)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    ext = "png" if fmt.upper() == "PNG" else "jpg"

    pdf_doc = fitz.open(pdf_path)
    total = len(pdf_doc)
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    results = []

    for i, page in enumerate(pdf_doc):
        pix = page.get_pixmap(matrix=mat)
        img_path = os.path.join(out_dir, f"{base_name}_page_{i + 1:03d}.{ext}")

        if fmt.upper() == "JPG":
            # 先保存为 PNG 再用 Pillow 转 JPG
            tmp_path = os.path.join(out_dir, f"_tmp_{i + 1}.png")
            pix.save(tmp_path)
            img = Image.open(tmp_path)
            if img.mode == "RGBA":
                img = img.convert("RGB")
            img.save(img_path, "JPEG", quality=95)
            os.remove(tmp_path)
        else:
            pix.save(img_path)

        results.append(img_path)

        if progress_cb:
            progress_cb(i + 1, total)

    pdf_doc.close()
    return results


def merge_pdfs(pdf_paths: list, output_path: str, progress_cb=None) -> str:
    """合并多个 PDF"""
    merged = fitz.open()
    total = len(pdf_paths)

    for i, pdf_path in enumerate(pdf_paths):
        src = fitz.open(pdf_path)
        merged.insert_pdf(src)
        src.close()

        if progress_cb:
            progress_cb(i + 1, total)

    merged.save(output_path)
    merged.close()
    return output_path


def split_pdf(pdf_path: str, pages_per_file: int, progress_cb=None) -> list:
    """按页数拆分 PDF"""
    out_dir = get_output_dir(pdf_path)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    results = []

    part_num = 1
    for start in range(0, total_pages, pages_per_file):
        end = min(start + pages_per_file, total_pages)
        out_path = os.path.join(out_dir, f"{base_name}_part_{part_num:02d}.pdf")

        new_doc = fitz.open()
        new_doc.insert_pdf(pdf_doc, from_page=start, to_page=end - 1)
        new_doc.save(out_path)
        new_doc.close()
        results.append(out_path)
        part_num += 1

        if progress_cb:
            progress_cb(end, total_pages)

    pdf_doc.close()
    return results


def images_to_pdf(image_paths: list, output_path: str, progress_cb=None) -> str:
    """多张图片合并为 PDF"""
    total = len(image_paths)
    pdf_doc = fitz.open()

    for i, img_path in enumerate(image_paths):
        img = Image.open(img_path)
        if img.mode == "RGBA":
            img = img.convert("RGB")
        # 使用 fitz 插入图片为页面
        img_w, img_h = img.size
        # 以图片尺寸创建页面（像素转点：1px = 1/72 inch * 72 = 1pt 近似）
        page = pdf_doc.new_page(width=img_w, height=img_h)
        page.insert_image(fitz.Rect(0, 0, img_w, img_h), filename=img_path)

        if progress_cb:
            progress_cb(i + 1, total)

    pdf_doc.save(output_path)
    pdf_doc.close()
    return output_path


def encrypt_pdf(pdf_path: str, password: str, progress_cb=None) -> str:
    """PDF 加密（设置打开密码）"""
    out_dir = get_output_dir(pdf_path)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    out_path = os.path.join(out_dir, f"{base_name}_encrypted.pdf")

    pdf_doc = fitz.open(pdf_path)
    perm = fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY | fitz.PDF_PERM_ANNOTATE
    pdf_doc.save(out_path, encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw=password,
                 user_pw=password, permissions=perm)
    pdf_doc.close()

    if progress_cb:
        progress_cb(1, 1)
    return out_path


def decrypt_pdf(pdf_path: str, password: str, progress_cb=None) -> str:
    """PDF 解密（移除密码）"""
    out_dir = get_output_dir(pdf_path)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    out_path = os.path.join(out_dir, f"{base_name}_decrypted.pdf")

    pdf_doc = fitz.open(pdf_path)
    if pdf_doc.needs_pass:
        if not pdf_doc.authenticate(password):
            pdf_doc.close()
            raise ValueError("密码错误，无法解密 PDF")
    pdf_doc.save(out_path)
    pdf_doc.close()

    if progress_cb:
        progress_cb(1, 1)
    return out_path


# ============================================================
#  GUI 界面
# ============================================================

class PDFToolApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF 万能转换工具")
        self.root.geometry("760x680")
        self.root.minsize(680, 600)

        # 风格设置
        style = ttk.Style()
        style.theme_use("clam")

        # 主容器
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=(10, 5))

        # 各功能选项卡
        self._build_tab_pdf_to_word()
        self._build_tab_pdf_to_images()
        self._build_tab_merge()
        self._build_tab_split()
        self._build_tab_images_to_pdf()
        self._build_tab_encrypt()
        self._build_tab_decrypt()

        # 底部进度条 + 日志
        bottom = ttk.Frame(root)
        bottom.pack(fill="x", padx=10, pady=(0, 10))

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(bottom, variable=self.progress_var,
                                            maximum=100, length=740)
        self.progress_bar.pack(fill="x", pady=(0, 5))

        self.log_text = scrolledtext.ScrolledText(bottom, height=6, font=("Consolas", 9),
                                                   state="disabled", wrap="word")
        self.log_text.pack(fill="x")

        # 运行状态
        self._running = False

    # ---------- 通用组件 ----------

    def _make_file_selector(self, parent, label_text, filetypes, mode="file"):
        """创建文件选择器行，返回 (frame, path_var, browse_func)"""
        frame = ttk.LabelFrame(parent, text=label_text, padding=10)
        path_var = tk.StringVar()

        entry = ttk.Entry(frame, textvariable=path_var)
        entry.pack(side="left", fill="x", expand=True, padx=(0, 5))

        def browse():
            if mode == "file":
                files = filedialog.askopenfilenames(filetypes=filetypes)
                if files:
                    path_var.set(" ; ".join(files))
            elif mode == "folder":
                folder = filedialog.askdirectory()
                if folder:
                    path_var.set(folder)

        btn = ttk.Button(frame, text="浏览...", command=browse)
        btn.pack(side="left")
        return frame, path_var, browse

    def _log(self, msg):
        """写入日志"""
        self.log_text.config(state="normal")
        self.log_text.insert("end", msg + "\n")
        self.log_text.see("end")
        self.log_text.config(state="disabled")
        self.root.update_idletasks()

    def _set_progress(self, current, total):
        """更新进度条"""
        pct = (current / total * 100) if total > 0 else 0
        self.progress_var.set(pct)
        self.root.update_idletasks()

    def _run_task(self, task_func, *args):
        """在后台线程运行任务"""
        if self._running:
            messagebox.showwarning("提示", "当前已有任务正在运行，请等待完成。")
            return

        self._running = True
        self.progress_var.set(0)

        def worker():
            try:
                task_func(*args)
                self._log("✅ 任务完成！")
            except Exception as e:
                self._log(f"❌ 错误：{e}")
                traceback.print_exc()
                messagebox.showerror("错误", str(e))
            finally:
                self._running = False
                self.progress_var.set(100)

        thread = threading.Thread(target=worker, daemon=True)
        thread.start()

    # ---------- 选项卡：PDF 转 Word ----------

    def _build_tab_pdf_to_word(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="📄 PDF转Word")

        ttk.Label(tab, text="选择 PDF 文件：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "PDF 文件",
                                                        [("PDF 文件", "*.pdf")])
        frame.pack(fill="x", pady=(0, 10))

        def do_convert():
            paths = path_var.get().split(" ; ")
            paths = [p.strip() for p in paths if p.strip()]
            if not paths:
                messagebox.showwarning("提示", "请先选择 PDF 文件！")
                return

            def task():
                for p in paths:
                    self._log(f"正在转换：{os.path.basename(p)} → Word ...")
                    out = pdf_to_word(p, progress_cb=self._set_progress)
                    self._log(f"  已保存：{out}")

            self._run_task(task)

        ttk.Button(tab, text="开始转换为 Word", command=do_convert).pack(pady=10)

    # ---------- 选项卡：PDF 转图片 ----------

    def _build_tab_pdf_to_images(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="🖼️ PDF转图片")

        ttk.Label(tab, text="选择 PDF 文件：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "PDF 文件",
                                                        [("PDF 文件", "*.pdf")])
        frame.pack(fill="x", pady=(0, 10))

        fmt_frame = ttk.Frame(tab)
        fmt_frame.pack(fill="x", pady=(0, 10))
        ttk.Label(fmt_frame, text="图片格式：").pack(side="left")
        fmt_var = tk.StringVar(value="PNG")
        ttk.Radiobutton(fmt_frame, text="PNG（无损）", variable=fmt_var, value="PNG").pack(side="left", padx=10)
        ttk.Radiobutton(fmt_frame, text="JPG（较小）", variable=fmt_var, value="JPG").pack(side="left")

        dpi_frame = ttk.Frame(tab)
        dpi_frame.pack(fill="x", pady=(0, 10))
        ttk.Label(dpi_frame, text="DPI 分辨率：").pack(side="left")
        dpi_var = tk.IntVar(value=200)
        for val, text in [(150, "150（标准）"), (200, "200（高清）"), (300, "300（印刷级）")]:
            ttk.Radiobutton(dpi_frame, text=text, variable=dpi_var, value=val).pack(side="left", padx=5)

        def do_convert():
            paths = path_var.get().split(" ; ")
            paths = [p.strip() for p in paths if p.strip()]
            if not paths:
                messagebox.showwarning("提示", "请先选择 PDF 文件！")
                return

            fmt = fmt_var.get()
            dpi = dpi_var.get()

            def task():
                for p in paths:
                    self._log(f"正在转换：{os.path.basename(p)} → {fmt} 图片 (DPI={dpi}) ...")
                    outs = pdf_to_images(p, fmt=fmt, dpi=dpi, progress_cb=self._set_progress)
                    self._log(f"  共生成 {len(outs)} 张图片，保存于：{os.path.dirname(outs[0]) if outs else 'N/A'}")

            self._run_task(task)

        ttk.Button(tab, text="开始转换为图片", command=do_convert).pack(pady=10)

    # ---------- 选项卡：合并 PDF ----------

    def _build_tab_merge(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="🔗 PDF合并")

        ttk.Label(tab, text="选择多个 PDF 文件（按选择顺序合并）：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "PDF 文件",
                                                        [("PDF 文件", "*.pdf")])
        frame.pack(fill="x", pady=(0, 10))

        ttk.Label(tab, text="合并后输出路径：").pack(anchor="w", pady=(0, 5))
        out_frame = ttk.Frame(tab)
        out_frame.pack(fill="x", pady=(0, 10))
        out_var = tk.StringVar()
        ttk.Entry(out_frame, textvariable=out_var).pack(side="left", fill="x", expand=True, padx=(0, 5))

        def browse_out():
            fp = filedialog.asksaveasfilename(defaultextension=".pdf",
                                               filetypes=[("PDF 文件", "*.pdf")],
                                               title="保存合并后的 PDF")
            if fp:
                out_var.set(fp)
        ttk.Button(out_frame, text="浏览...", command=browse_out).pack(side="left")

        def do_merge():
            paths = path_var.get().split(" ; ")
            paths = [p.strip() for p in paths if p.strip()]
            if len(paths) < 2:
                messagebox.showwarning("提示", "请至少选择 2 个 PDF 文件！")
                return
            out_path = out_var.get().strip()
            if not out_path:
                messagebox.showwarning("提示", "请指定输出路径！")
                return

            def task():
                self._log(f"正在合并 {len(paths)} 个 PDF 文件 ...")
                merge_pdfs(paths, out_path, progress_cb=self._set_progress)
                self._log(f"  合并完成：{out_path}")

            self._run_task(task)

        ttk.Button(tab, text="开始合并", command=do_merge).pack(pady=10)

    # ---------- 选项卡：拆分 PDF ----------

    def _build_tab_split(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="✂️ PDF拆分")

        ttk.Label(tab, text="选择 PDF 文件：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "PDF 文件",
                                                        [("PDF 文件", "*.pdf")])
        frame.pack(fill="x", pady=(0, 10))

        pages_frame = ttk.Frame(tab)
        pages_frame.pack(fill="x", pady=(0, 10))
        ttk.Label(pages_frame, text="每个文件包含页数：").pack(side="left")
        pages_var = tk.IntVar(value=1)
        spin = ttk.Spinbox(pages_frame, from_=1, to=999, textvariable=pages_var, width=8)
        spin.pack(side="left", padx=10)

        def do_split():
            path = path_var.get().strip()
            if not path:
                messagebox.showwarning("提示", "请先选择 PDF 文件！")
                return
            pages = pages_var.get()
            if pages < 1:
                messagebox.showwarning("提示", "每个文件至少包含 1 页！")
                return

            def task():
                self._log(f"正在拆分：{os.path.basename(path)}，每 {pages} 页一份 ...")
                outs = split_pdf(path, pages, progress_cb=self._set_progress)
                self._log(f"  共拆分为 {len(outs)} 个文件，保存于：{os.path.dirname(outs[0]) if outs else 'N/A'}")

            self._run_task(task)

        ttk.Button(tab, text="开始拆分", command=do_split).pack(pady=10)

    # ---------- 选项卡：图片转 PDF ----------

    def _build_tab_images_to_pdf(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="📷 图片转PDF")

        ttk.Label(tab, text="选择图片文件（支持 PNG/JPG/BMP/GIF/TIFF）：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "图片文件",
                                                        [("图片文件", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff")])
        frame.pack(fill="x", pady=(0, 10))

        ttk.Label(tab, text="输出 PDF 路径：").pack(anchor="w", pady=(0, 5))
        out_frame = ttk.Frame(tab)
        out_frame.pack(fill="x", pady=(0, 10))
        out_var = tk.StringVar()
        ttk.Entry(out_frame, textvariable=out_var).pack(side="left", fill="x", expand=True, padx=(0, 5))

        def browse_out():
            fp = filedialog.asksaveasfilename(defaultextension=".pdf",
                                               filetypes=[("PDF 文件", "*.pdf")],
                                               title="保存 PDF")
            if fp:
                out_var.set(fp)
        ttk.Button(out_frame, text="浏览...", command=browse_out).pack(side="left")

        def do_convert():
            paths = path_var.get().split(" ; ")
            paths = [p.strip() for p in paths if p.strip()]
            if not paths:
                messagebox.showwarning("提示", "请先选择图片文件！")
                return
            out_path = out_var.get().strip()
            if not out_path:
                messagebox.showwarning("提示", "请指定输出 PDF 路径！")
                return

            def task():
                self._log(f"正在将 {len(paths)} 张图片合并为 PDF ...")
                images_to_pdf(paths, out_path, progress_cb=self._set_progress)
                self._log(f"  转换完成：{out_path}")

            self._run_task(task)

        ttk.Button(tab, text="开始转换", command=do_convert).pack(pady=10)

    # ---------- 选项卡：加密 PDF ----------

    def _build_tab_encrypt(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="🔒 PDF加密")

        ttk.Label(tab, text="选择 PDF 文件：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "PDF 文件",
                                                        [("PDF 文件", "*.pdf")])
        frame.pack(fill="x", pady=(0, 10))

        ttk.Label(tab, text="设置打开密码：").pack(anchor="w", pady=(0, 5))
        pwd_frame = ttk.Frame(tab)
        pwd_frame.pack(fill="x", pady=(0, 10))
        pwd_var = tk.StringVar()
        ttk.Entry(pwd_frame, textvariable=pwd_var, show="*").pack(side="left", fill="x", expand=True, padx=(0, 5))
        ttk.Label(pwd_frame, text="确认密码：").pack(side="left", padx=(10, 5))
        pwd2_var = tk.StringVar()
        ttk.Entry(pwd_frame, textvariable=pwd2_var, show="*").pack(side="left")

        def do_encrypt():
            path = path_var.get().strip()
            if not path:
                messagebox.showwarning("提示", "请先选择 PDF 文件！")
                return
            pwd = pwd_var.get()
            pwd2 = pwd2_var.get()
            if not pwd:
                messagebox.showwarning("提示", "请输入密码！")
                return
            if pwd != pwd2:
                messagebox.showwarning("提示", "两次输入的密码不一致！")
                return

            def task():
                self._log(f"正在加密：{os.path.basename(path)} ...")
                out = encrypt_pdf(path, pwd, progress_cb=self._set_progress)
                self._log(f"  加密完成：{out}")

            self._run_task(task)

        ttk.Button(tab, text="开始加密", command=do_encrypt).pack(pady=10)

    # ---------- 选项卡：解密 PDF ----------

    def _build_tab_decrypt(self):
        tab = ttk.Frame(self.notebook, padding=15)
        self.notebook.add(tab, text="🔓 PDF解密")

        ttk.Label(tab, text="选择加密的 PDF 文件：").pack(anchor="w", pady=(0, 5))
        frame, path_var, _ = self._make_file_selector(tab, "PDF 文件",
                                                        [("PDF 文件", "*.pdf")])
        frame.pack(fill="x", pady=(0, 10))

        ttk.Label(tab, text="输入打开密码：").pack(anchor="w", pady=(0, 5))
        pwd_frame = ttk.Frame(tab)
        pwd_frame.pack(fill="x", pady=(0, 10))
        pwd_var = tk.StringVar()
        ttk.Entry(pwd_frame, textvariable=pwd_var, show="*").pack(side="left", fill="x", expand=True)

        def do_decrypt():
            path = path_var.get().strip()
            if not path:
                messagebox.showwarning("提示", "请先选择 PDF 文件！")
                return
            pwd = pwd_var.get()
            if not pwd:
                messagebox.showwarning("提示", "请输入密码！")
                return

            def task():
                self._log(f"正在解密：{os.path.basename(path)} ...")
                out = decrypt_pdf(path, pwd, progress_cb=self._set_progress)
                self._log(f"  解密完成：{out}")

            self._run_task(task)

        ttk.Button(tab, text="开始解密", command=do_decrypt).pack(pady=10)


# ============================================================
#  入口
# ============================================================

def main():
    root = tk.Tk()
    app = PDFToolApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
