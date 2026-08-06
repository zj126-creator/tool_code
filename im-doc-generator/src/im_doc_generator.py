# -*- coding: utf-8 -*-
"""
IM在线文档表生成器 v2
直接基于模板XML克隆，确保格式与模板完全一致
"""

import os
import sys
import re
from datetime import datetime
from copy import deepcopy

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


def read_excel_data(excel_path, selected_risks=None):
    """读取日计划导出Excel，返回作业列表"""
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    jobs = []
    for row in ws.iter_rows(min_row=3, max_row=ws.max_row, values_only=False):
        plan_code = row[0].value or ''
        plan_name = row[1].value or ''
        voltage = row[5].value or ''
        job_type = row[6].value or ''
        job_content = row[8].value or ''
        risk_level = row[14].value or ''
        is_live = row[16].value or ''
        plan_status = row[24].value or ''

        # 工作负责人 (AV=47姓名, AX=49单位) 不取电话AW=48
        responsible_name = row[47].value or '' if len(row) > 47 else ''
        responsible_unit = row[49].value or '' if len(row) > 49 else ''

        # 到岗到位人员 (BM=64姓名, BN=65单位) 不取电话BO=66
        arrival_names = row[64].value or '' if len(row) > 64 else ''
        arrival_units = row[65].value or '' if len(row) > 65 else ''

        # 安全督查人员 (BT=71姓名, BU=72单位) 不取电话BV=73
        inspector_names = row[71].value or '' if len(row) > 71 else ''
        inspector_units = row[72].value or '' if len(row) > 72 else ''

        if selected_risks and risk_level not in selected_risks:
            continue

        # 构建现场描述
        site_name = plan_name if plan_name else job_content
        if job_content:
            job_content = str(job_content).replace('\n', '；').strip().rstrip('；')

        desc_parts = []
        if site_name:
            desc_parts.append(site_name)
        if job_content and job_content != site_name:
            desc_parts.append(f'工作内容：{job_content}')
        if risk_level:
            desc_parts.append(f'{risk_level}风险')
        if is_live and str(is_live) == '是':
            desc_parts.append('带电作业')

        site_desc = '。'.join(desc_parts) + '。' if desc_parts else ''

        # 工作负责人（不含电话）
        responsible_info = ''
        if responsible_name:
            responsible_info = responsible_name
            if responsible_unit:
                responsible_info = f'{responsible_name}（{responsible_unit}）'

        # 同进同出人员
        arrival_info = _format_personnel(arrival_names, arrival_units)

        # 安全检查人员
        inspector_info = _format_personnel(inspector_names, inspector_units)

        jobs.append({
            'site_desc': site_desc,
            'risk_level': str(risk_level),
            'responsible': responsible_info,
            'arrival': arrival_info,
            'inspector': inspector_info,
        })

    return jobs


def _format_personnel(names_str, units_str):
    """格式化人员信息"""
    if not names_str:
        return ''
    names = [n.strip() for n in str(names_str).replace('，', '\n').replace(',', '\n').split('\n') if n.strip()]
    if not names:
        return ''
    if units_str:
        units = [u.strip() for u in str(units_str).replace('，', '\n').replace(',', '\n').split('\n') if u.strip()]
        pairs = []
        for i, n in enumerate(names):
            u = units[i] if i < len(units) else ''
            pairs.append(f'{n}（{u}）' if u else n)
        return '，'.join(pairs)
    return '，'.join(names)


def get_risk_levels(excel_path):
    """从Excel中获取所有风险等级选项"""
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active
    risks = set()
    for row in ws.iter_rows(min_row=3, max_row=ws.max_row, min_col=15, max_col=15):
        for cell in row:
            if cell.value:
                risks.add(str(cell.value))
    return sorted(risks)


def _set_run_text(run, text):
    """设置run的文本（保留run格式）"""
    run.text = text


def _set_cell_text_preserve_format(cell, text, alignment=None):
    """设置单元格文本，保留原有格式"""
    # 清除所有段落内容但保留第一个段落格式
    for i in range(len(cell.paragraphs) - 1, 0, -1):
        cell.paragraphs[i]._p.getparent().remove(cell.paragraphs[i]._p)
    
    p = cell.paragraphs[0]
    # 清除run但保留段落格式
    for run in p.runs:
        run._r.getparent().remove(run._r)
    
    # 设置对齐
    if alignment is not None:
        p.alignment = alignment
    
    if text:
        lines = str(text).split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                run = p.add_run()
                br = OxmlElement('w:br')
                run._r.append(br)
            run = p.add_run(line)
            # 复制模板字体
            run.font.name = '方正仿宋_GBK'
            run.font.size = Pt(10.5)
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
            rFonts.set(qn('w:ascii'), '方正仿宋_GBK')
            rFonts.set(qn('w:hAnsi'), '方正仿宋_GBK')
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '21')


def generate_word(jobs, template_path, output_path):
    """生成Word文档 - 基于模板XML克隆"""
    # 读取模板
    template_doc = Document(template_path)
    
    today = datetime.now()
    month_day = f'{today.month:02d}月{today.day:02d}日'
    
    # 修改标题日期
    for p in template_doc.paragraphs:
        if '作业现场安全管控人员安排表' in p.text:
            # 找到包含标题文字的run
            for run in p.runs:
                if '作业现场安全管控人员安排表' in run.text:
                    run.text = f'作业现场安全管控人员安排表（{month_day}）'
            break
    
    # 获取模板表格
    template_table = template_doc.tables[0]
    template_tbl = template_table._tbl
    
    # 模板中每个作业现场是5行，模板有5个作业现场(25行)
    # 我们需要根据实际jobs数量调整
    
    # 先获取模板中第一个作业现场的5行作为模板行
    template_rows = []
    for ri in range(5):
        tr = template_tbl.findall(qn('w:tr'))[ri]
        template_rows.append(deepcopy(tr))
    
    # 删除模板表格中所有行
    all_trs = template_tbl.findall(qn('w:tr'))
    for tr in all_trs:
        template_tbl.remove(tr)
    
    # 为每个job生成5行
    for job_idx, job in enumerate(jobs):
        for template_tr in template_rows:
            new_tr = deepcopy(template_tr)
            template_tbl.append(new_tr)
        
        # 获取刚添加的5行
        all_trs = template_tbl.findall(qn('w:tr'))
        base = job_idx * 5
        
        # 第1行: 作业现场 + 工作负责人
        row0 = all_trs[base]
        cells0 = row0.findall(qn('w:tc'))
        _update_cell_xml(cells0[0], f'{job_idx+1}.作业现场名称及作业内容、风险等级：\n{job["site_desc"]}', is_site_desc=True)
        # 工作负责人标签cell
        _update_cell_xml(cells0[1], '工作负责人及岗位', is_label=True)
        # 工作负责人内容cell
        _update_cell_xml(cells0[2], job['responsible'])
        
        # 第2行: 同进同出
        row1 = all_trs[base + 1]
        cells1 = row1.findall(qn('w:tc'))
        _update_cell_xml(cells1[0], f'{job_idx+1}.作业现场名称及作业内容、风险等级：\n{job["site_desc"]}', is_site_desc=True)
        _update_cell_xml(cells1[1], '同进同出人员及岗位', is_label=True)
        _update_cell_xml(cells1[2], job['arrival'])
        
        # 第3行: 安全检查
        row2 = all_trs[base + 2]
        cells2 = row2.findall(qn('w:tc'))
        _update_cell_xml(cells2[0], f'{job_idx+1}.作业现场名称及作业内容、风险等级：\n{job["site_desc"]}', is_site_desc=True)
        _update_cell_xml(cells2[1], '安全检查人员及岗位', is_label=True)
        _update_cell_xml(cells2[2], job['inspector'])
        
        # 从第二个作业现场开始，C0对齐改为left（模板中第一个是center，后续是left）
        if job_idx > 0:
            for ridx in range(base, base + 3):
                tc0 = all_trs[ridx].findall(qn('w:tc'))[0]
                p0 = tc0.findall(qn('w:p'))[0]
                pPr = p0.find(qn('w:pPr'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        jc.set(qn('w:val'), 'left')
        
        # 第4行: 监护表头 - 保持不变（序号/监护地点及内容/监护要点/监护人及岗位）
        # 不需要修改，模板内容已经是正确的
        
        # 第5行: 空数据行 - 保持不变
        # 不需要修改
    
    template_doc.save(output_path)
    return output_path


def _make_run(text, bold=False):
    """创建一个带完整格式的run元素（与模板一致）"""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rFonts.set(qn('w:ascii'), '方正仿宋_GBK')
    rFonts.set(qn('w:hAnsi'), '方正仿宋_GBK')
    rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
    rFonts.set(qn('w:cs'), '方正仿宋_GBK')
    rPr.append(rFonts)
    
    b = OxmlElement('w:b')
    if not bold:
        b.set(qn('w:val'), '0')
    rPr.append(b)
    
    bCs = OxmlElement('w:bCs')
    if not bold:
        bCs.set(qn('w:val'), '0')
    rPr.append(bCs)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'auto')
    rPr.append(color)
    
    kern = OxmlElement('w:kern')
    kern.set(qn('w:val'), '2')
    rPr.append(kern)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'none')
    rPr.append(highlight)
    
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def _make_break_run():
    """创建一个换行run"""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rFonts.set(qn('w:ascii'), '方正仿宋_GBK')
    rFonts.set(qn('w:hAnsi'), '方正仿宋_GBK')
    rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
    rFonts.set(qn('w:cs'), '方正仿宋_GBK')
    rPr.append(rFonts)
    b = OxmlElement('w:b')
    b.set(qn('w:val'), '0')
    rPr.append(b)
    bCs = OxmlElement('w:bCs')
    bCs.set(qn('w:val'), '0')
    rPr.append(bCs)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'auto')
    rPr.append(color)
    kern = OxmlElement('w:kern')
    kern.set(qn('w:val'), '2')
    rPr.append(kern)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'none')
    rPr.append(highlight)
    r.append(rPr)
    br = OxmlElement('w:br')
    r.append(br)
    return r


# 模板中作业现场描述的固定前缀
_SITE_PREFIX = '作业现场名称及作业内容、风险等级'


def _update_cell_xml(tc, text, is_site_desc=False, is_label=False):
    """更新单元格tc的XML中的文本内容，保留格式
    
    is_site_desc=True时，将"作业现场名称及作业内容、风险等级"部分加粗（与模板一致）
    is_label=True时，整个文本加粗（用于标签单元格如"工作负责人及岗位"）
    """
    paragraphs = tc.findall(qn('w:p'))
    if not paragraphs:
        return
    
    # 保留第一个段落，删除其余
    for p in paragraphs[1:]:
        tc.remove(p)
    
    p = paragraphs[0]
    
    # 删除所有run
    for r in p.findall(qn('w:r')):
        p.remove(r)
    
    if not text:
        return
    
    if is_site_desc:
        # 模板格式：run[0]=序号(不粗), run[1]=固定前缀(粗体), run[2]=冒号+内容(不粗)
        # 解析 text 格式: "N.作业现场名称及作业内容、风险等级：\n具体内容"
        import re
        m = re.match(r'^(\d+\.)' + re.escape(_SITE_PREFIX) + r'[：:]([\s\S]*)$', text)
        if m:
            num = m.group(1)
            content = m.group(2)
            # 第一行: 序号(不粗) + 固定前缀(粗体) + 冒号及内容(不粗)
            # 冒号归到内容部分
            if content.startswith('\n'):
                content = content[1:]
            
            # 检查内容中是否有换行（模板中有些是单行，有些有多行）
            content_lines = content.split('\n')
            
            # run[0]: 序号 (不粗)
            p.append(_make_run(num, bold=False))
            # run[1]: 固定前缀 (粗体)
            p.append(_make_run(_SITE_PREFIX, bold=True))
            # run[2]: 冒号 + 第一行内容 (不粗)
            if content_lines[0]:
                p.append(_make_run('：' + content_lines[0], bold=False))
            else:
                p.append(_make_run('：', bold=False))
            
            # 后续行: 换行 + 内容 (不粗)
            for line in content_lines[1:]:
                p.append(_make_break_run())
                p.append(_make_run(line, bold=False))
        else:
            # 格式不匹配，回退到普通方式
            lines = str(text).split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    p.append(_make_break_run())
                p.append(_make_run(line, bold=False))
    else:
        bold = is_label
        lines = str(text).split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                p.append(_make_break_run())
            p.append(_make_run(line, bold=bold))


# ==================== GUI ====================
def create_gui():
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox

    root = tk.Tk()
    root.title('IM在线文档表生成器')
    root.geometry('600x520')
    root.resizable(False, False)

    style = ttk.Style()
    style.configure('Title.TLabel', font=('微软雅黑', 14, 'bold'))
    style.configure('Sub.TLabel', font=('微软雅黑', 10))
    style.configure('Btn.TButton', font=('微软雅黑', 10))

    ttk.Label(root, text='IM在线文档表生成器', style='Title.TLabel').pack(pady=10)

    file_frame = ttk.LabelFrame(root, text='文件选择', padding=10)
    file_frame.pack(fill='x', padx=20, pady=5)

    # Excel
    ttk.Label(file_frame, text='日计划导出(Excel)：', style='Sub.TLabel').grid(row=0, column=0, sticky='w', pady=3)
    excel_var = tk.StringVar()
    ttk.Entry(file_frame, textvariable=excel_var, width=45).grid(row=0, column=1, pady=3, padx=5)

    def browse_excel():
        path = filedialog.askopenfilename(
            title='选择日计划导出Excel文件',
            filetypes=[('Excel文件', '*.xlsx *.xls'), ('所有文件', '*.*')]
        )
        if path:
            excel_var.set(path)
            try:
                risks = get_risk_levels(path)
                for widget in risk_frame.winfo_children():
                    if not isinstance(widget, ttk.Frame):
                        widget.destroy()
                risk_vars.clear()
                for r in risks:
                    var = tk.BooleanVar(value=True)
                    risk_vars[r] = var
                    ttk.Checkbutton(risk_frame, text=r, variable=var).pack(side='left', padx=5)
            except Exception as e:
                messagebox.showerror('错误', f'读取Excel失败：{e}')

    ttk.Button(file_frame, text='浏览...', style='Btn.TButton', command=browse_excel).grid(row=0, column=2, pady=3)

    # Word模板
    ttk.Label(file_frame, text='Word模板文件：', style='Sub.TLabel').grid(row=1, column=0, sticky='w', pady=3)
    template_var = tk.StringVar()
    ttk.Entry(file_frame, textvariable=template_var, width=45).grid(row=1, column=1, pady=3, padx=5)

    def browse_template():
        path = filedialog.askopenfilename(
            title='选择Word模板文件',
            filetypes=[('Word文档', '*.docx'), ('所有文件', '*.*')]
        )
        if path:
            template_var.set(path)

    ttk.Button(file_frame, text='浏览...', style='Btn.TButton', command=browse_template).grid(row=1, column=2, pady=3)

    # 输出目录
    ttk.Label(file_frame, text='输出目录：', style='Sub.TLabel').grid(row=2, column=0, sticky='w', pady=3)
    output_var = tk.StringVar()
    ttk.Entry(file_frame, textvariable=output_var, width=45).grid(row=2, column=1, pady=3, padx=5)

    def browse_output():
        path = filedialog.askdirectory(title='选择输出目录')
        if path:
            output_var.set(path)

    ttk.Button(file_frame, text='浏览...', style='Btn.TButton', command=browse_output).grid(row=2, column=2, pady=3)

    # 风险等级
    risk_frame = ttk.LabelFrame(root, text='风险等级筛选（可多选）', padding=10)
    risk_frame.pack(fill='x', padx=20, pady=5)
    risk_vars = {}
    ttk.Label(risk_frame, text='请先选择Excel文件后自动加载风险等级', style='Sub.TLabel').pack()

    btn_frame = ttk.Frame(risk_frame)
    btn_frame.pack(fill='x', pady=3)

    def select_all():
        for var in risk_vars.values():
            var.set(True)

    def deselect_all():
        for var in risk_vars.values():
            var.set(False)

    ttk.Button(btn_frame, text='全选', command=select_all).pack(side='left', padx=5)
    ttk.Button(btn_frame, text='全不选', command=deselect_all).pack(side='left', padx=5)

    def generate():
        excel_path = excel_var.get()
        template_path = template_var.get()
        output_dir = output_var.get()

        if not excel_path or not os.path.exists(excel_path):
            messagebox.showerror('错误', '请选择有效的Excel文件')
            return
        if not template_path or not os.path.exists(template_path):
            messagebox.showerror('错误', '请选择有效的Word模板文件')
            return
        if not output_dir:
            messagebox.showerror('错误', '请选择输出目录')
            return

        selected_risks = [r for r, v in risk_vars.items() if v.get()]
        if not selected_risks:
            messagebox.showerror('错误', '请至少选择一个风险等级')
            return

        try:
            jobs = read_excel_data(excel_path, selected_risks)
            if not jobs:
                messagebox.showwarning('提示', '没有符合条件的作业计划')
                return

            today = datetime.now()
            output_name = f'IM在线文档表-{today.month:02d}{today.day:02d}.docx'
            output_path = os.path.join(output_dir, output_name)
            generate_word(jobs, template_path, output_path)

            messagebox.showinfo('成功', f'已生成：{output_path}\n共{len(jobs)}条作业计划')
        except Exception as e:
            messagebox.showerror('错误', f'生成失败：{e}')

    ttk.Button(root, text='生成Word文档', style='Btn.TButton', command=generate).pack(pady=15)
    ttk.Label(root, text='v2.0 | 从日计划导出Excel生成IM在线文档表', style='Sub.TLabel').pack(side='bottom', pady=5)

    root.mainloop()


if __name__ == '__main__':
    create_gui()
